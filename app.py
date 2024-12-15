from flask import Flask, render_template, request
import os
import string
from docx import Document
import io

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
ALLOWED_EXTENSIONS = {'txt', 'docx'}

main_text = []
main_gui = []
source_texts = []
source_gui = []
plag_ratio = 0.0
file_dict = {}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def read_file(file_content, file_type):
    main_text.clear()
    main_gui.clear()
    if file_type == 'txt':
        text_lines = file_content.splitlines()
        for line in text_lines:
            words = line.split()
            for word in words:
                cleaned_word = word.translate(str.maketrans('', '', string.punctuation)).lower()
                main_text.append(cleaned_word)
                main_gui.append({
                    'word': word,
                    'is_plagiarized': False,
                    'source_info': None,
                    'is_new_line': False
                })
            if words:
                main_gui[-1]['is_new_line'] = True
    elif file_type == 'docx':
        doc = Document(io.BytesIO(file_content))
        for paragraph in doc.paragraphs:
            words = paragraph.text.split()
            for word in words:
                cleaned_word = word.translate(str.maketrans('', '', string.punctuation)).lower()
                main_text.append(cleaned_word)
                main_gui.append({
                    'word': word,
                    'is_plagiarized': False,
                    'source_info': None,
                    'is_new_line': False
                })
            if words:
                main_gui[-1]['is_new_line'] = True

def read_files(file_contents, file_types):
    source_texts.clear()
    source_gui.clear()
    for content, file_type in zip(file_contents, file_types):
        if file_type == 'txt':
            text_lines = content.splitlines()
            file_words = []
            file_gui = []
            for line in text_lines:
                words = line.split()
                for word in words:
                    cleaned_word = word.translate(str.maketrans('', '', string.punctuation)).lower()
                    file_words.append(cleaned_word)
                    file_gui.append({
                        'word': word,
                        'is_plagiarized': False,
                        'is_new_line': False
                    })
                if words:
                    file_gui[-1]['is_new_line'] = True
            source_texts.append(file_words)
            source_gui.append(file_gui)
        elif file_type == 'docx':
            doc = Document(io.BytesIO(content))
            file_words = []
            file_gui = []
            for paragraph in doc.paragraphs:
                words = paragraph.text.split()
                for word in words:
                    cleaned_word = word.translate(str.maketrans('', '', string.punctuation)).lower()
                    file_words.append(cleaned_word)
                    file_gui.append({
                        'word': word,
                        'is_plagiarized': False,
                        'is_new_line': False
                    })
                if words:
                    file_gui[-1]['is_new_line'] = True
            source_texts.append(file_words)
            source_gui.append(file_gui)

def delete_gui_marks():
    for word in main_gui:
        word['is_plagiarized'] = False
        word['source_info'] = None
    for file_gui in source_gui:
        for word in file_gui:
            word['is_plagiarized'] = False

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['GET', 'POST'])
def upload():
    if request.method == 'POST':
        N = int(request.form.get('N'))
        main_file = request.files.get('mainfile')
        other_files = request.files.getlist('otherfiles')
        
        if not main_file or not other_files:
            return render_template('index.html', error='No files selected.')
        
        if main_file and allowed_file(main_file.filename):
            main_file_type = main_file.filename.rsplit('.', 1)[1].lower()
            file_content = main_file.read()
            file_dict[0] = main_file.filename
            read_file(file_content, main_file_type)
        else:
            return render_template('index.html', error='Invalid main file type.')
        
        file_contents = []
        file_types = []
        for i, file in enumerate(other_files, start=1):
            if file and allowed_file(file.filename):
                file_type = file.filename.rsplit('.', 1)[1].lower()
                file_contents.append(file.read())
                file_types.append(file_type)
                file_dict[i] = file.filename
            else:
                return render_template('index.html', error=f'Invalid file type for file {i}.')
        
        read_files(file_contents, file_types)
        result = check_plagiarism(N)
        return render_template('result.html', result=result, plag_ratio=result['plag_ratio'])
    
    return render_template('index.html')

def check_plagiarism(N):
    delete_gui_marks()
    n = N - 1
    main_window = main_text[:N-1].copy()
    main_len = len(main_text)
    plag_count = 0
    word_count = (main_len - n) * N
    
    for id in range(main_len - n):
        main_window.append(main_text[id + n])
        for file_id in range(len(source_texts)):
            for j in range(len(source_texts[file_id]) - n):
                if main_window == source_texts[file_id][j:j+N]:
                    for k in range(N):
                        word_index = id + k
                        source_word_index = j + k
                        if word_index < len(main_gui) and source_word_index < len(source_gui[file_id]):
                            main_gui[word_index]['is_plagiarized'] = True
                            main_gui[word_index]['source_info'] = [file_id, source_word_index]
                            source_gui[file_id][source_word_index]['is_plagiarized'] = True
                            plag_count += 1
        main_window.pop(0)
    
    if word_count == 0:
        plag_ratio = 0.0
    else:
        plag_ratio = (plag_count / word_count) * 100
    
    return {
        'main_html': generate_html(-1),
        'sources_html': [generate_html(file_id) for file_id in range(1, len(file_dict))],
        'plag_ratio': plag_ratio
    }

def generate_html(file_id):
    if file_id == -1:
        main_html = []
        main_html.append("<p>")
        for word in main_gui:
            if not word['is_plagiarized']:
                main_html.append(word['word'] + " ")
            else:
                link = f'<a href="#w{word["source_info"][0]}-{word["source_info"][1]}" class="plagiarized-link">{word["word"]}</a> '
                main_html.append(link)
            if word['is_new_line']:
                main_html.append("</p><p>")
        main_html.append("</p>")
        return "".join(main_html)
    else:
        source_html = []
        source_html.append(f"<h2>Source {file_id} - {file_dict[file_id]}</h2><p>")
        for word in source_gui[file_id - 1]:
            if not word['is_plagiarized']:
                source_html.append(word['word'] + " ")
            else:
                source_html.append(f'<span class="plagiarized-word">{word["word"]}</span> ')
            if word['is_new_line']:
                source_html.append("</p><p>")
        source_html.append("</p>")
        return "".join(source_html)

if __name__ == '__main__':
    app.run(debug=True)
